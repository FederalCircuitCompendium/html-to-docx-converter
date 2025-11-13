import streamlit as st
from io import BytesIO
from pathlib import Path
import html as py_html
import tempfile
from typing import Optional
import re

from docx import Document

APP_DIR = Path(__file__).parent
REF_DOCX = APP_DIR / "assets" / "reference.docx"

st.set_page_config(page_title="HTML → DOCX", page_icon="📄", layout="wide")
st.title("HTML → DOCX (templated)")

st.markdown(
    "Use this form to convert HTML into a DOCX file. "
    "Fields marked with * are required."
)

# --- Inputs (in a single form for predictable keyboard behavior) ---
with st.form("converter_form"):
    title = st.text_input(
        "Document title (optional)",
        value="",
        help="Used as the Word document title. If left blank, a default title is used."
    )

    html_body = st.text_area(
        "HTML source *",
        value="",
        height=420,
        help="Paste the raw HTML you want to convert. This field is required."
    )

    start_level = st.number_input(
        "Heading level for first body heading",
        min_value=1,
        max_value=9,
        value=1,
        step=1,
        help=(
            "Use 1 to keep existing heading levels. "
            "Higher numbers demote all headings "
            "(for example, 3 makes the first heading “Heading 3”)."
        )
    )

    strong_emph = st.checkbox(
        "Map bold/italic to Word Strong/Emphasis styles",
        value=True,
        help="When checked, bold text uses the Strong style and italic text uses Emphasis."
    )

    col1, col2 = st.columns([1, 1])
    with col1:
        st.markdown(
            "- The document title is stored as metadata.\n"
            "- The document language is set to `en-US`.\n"
            f"- Template reference file: {'found' if REF_DOCX.exists() else 'not found (conversion still works)'}"
        )

    submitted = st.form_submit_button("Convert HTML to DOCX")
    
def add_page_numbers(doc: Document):
    """
    Add centered page numbers to the footer of each section.
    Uses a PAGE field so Word can renumber automatically.
    """
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    for section in doc.sections:
        footer = section.footer

        # Reuse first paragraph if it's effectively empty; otherwise add a new one.
        if footer.paragraphs:
            paragraph = footer.paragraphs[0]
            if paragraph.text.strip():
                paragraph = footer.add_paragraph()
        else:
            paragraph = footer.add_paragraph()

        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add the PAGE field: { PAGE }
        run = paragraph.add_run()

        fld_begin = OxmlElement("w:fldChar")
        fld_begin.set(qn("w:fldCharType"), "begin")

        instr_text = OxmlElement("w:instrText")
        instr_text.set(qn("xml:space"), "preserve")
        instr_text.text = " PAGE "

        fld_end = OxmlElement("w:fldChar")
        fld_end.set(qn("w:fldCharType"), "end")

        run._r.append(fld_begin)
        run._r.append(instr_text)
        run._r.append(fld_end)


def apply_language_en_us(doc: Document):
    ns = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"

    def tag_runs(container):
        for p in container.paragraphs:
            for r in p.runs:
                if r._element.rPr is None:
                    r._element._add_rPr()
                r._element.rPr.set(ns + "lang", "en-US")

    tag_runs(doc)
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                tag_runs(cell)


def remap_headings(doc: Document, start_level: int):
    """
    Remap all paragraph heading styles so that original Heading 1 becomes
    Heading <start_level>, Heading 2 -> start_level+1, etc.

    This is robust to style names like 'Heading 2 Char' by extracting the
    first integer from the style name.
    """
    delta = max(0, start_level - 1)
    if delta == 0:
        return

    for p in doc.paragraphs:
        s = p.style
        if not s:
            continue

        nm = getattr(s, "name", str(s))
        if not nm.startswith("Heading"):
            continue

        # Find first integer in the style name (e.g. 'Heading 2 Char' -> 2)
        m = re.search(r"(\d+)", nm)
        if not m:
            continue

        lvl = int(m.group(1))
        new_lvl = min(9, lvl + delta)
        target_style_name = f"Heading {new_lvl}"

        try:
            p.style = target_style_name
        except KeyError:
            # If the template doesn't define that heading level, just leave it
            pass


def normalize_heading_levels(doc: Document):
    """
    Normalize heading levels so that the *smallest* existing heading level
    becomes Heading 1, and higher levels are shifted up accordingly.

    Example:
      If the template causes <h1>, <h2>, <h3> to map to Heading 2, 3, 4:
      - min level = 2
      - We shift:
          Heading 2 -> Heading 1
          Heading 3 -> Heading 2
          Heading 4 -> Heading 3
    """
    heading_levels = []

    # First pass: collect all heading levels in use
    for p in doc.paragraphs:
        s = p.style
        if not s:
            continue
        nm = getattr(s, "name", "")
        if not nm.startswith("Heading"):
            continue
        m = re.search(r"(\d+)", nm)
        if not m:
            continue
        heading_levels.append(int(m.group(1)))

    if not heading_levels:
        return  # no headings to normalize

    min_lvl = min(heading_levels)
    if min_lvl <= 1:
        # Already starting at Heading 1 or below; nothing to do.
        return

    delta = min_lvl - 1  # how much we need to shift up

    # Second pass: shift all headings >= min_lvl up by delta
    for p in doc.paragraphs:
        s = p.style
        if not s:
            continue
        nm = getattr(s, "name", "")
        if not nm.startswith("Heading"):
            continue
        m = re.search(r"(\d+)", nm)
        if not m:
            continue
        lvl = int(m.group(1))
        if lvl < min_lvl:
            continue  # leave anything "below" the first heading alone

        new_lvl = max(1, lvl - delta)
        target_style_name = f"Heading {min(new_lvl, 9)}"
        try:
            p.style = target_style_name
        except KeyError:
            # If that heading level doesn't exist in the template, skip
            pass


def bold_italic_to_character_styles(doc: Document, use=True):
    if not use:
        return
    for p in doc.paragraphs:
        for r in p.runs:
            if r.bold:
                try:
                    r.style = "Strong"
                except Exception:
                    pass
            if r.italic:
                try:
                    r.style = "Emphasis"
                except Exception:
                    pass


def try_pandoc_convert(html_str: str, title: str, out_path: Path, reference: Optional[Path]):
    try:
        import pypandoc
        try:
            pypandoc.get_pandoc_version()
        except OSError:
            pypandoc.download_pandoc()
        extra_args = ["--metadata=lang=en-US", f"--metadata=title={title}"]
        if reference and reference.exists():
            extra_args.append(f"--reference-doc={reference}")
        pypandoc.convert_text(
            html_str,
            to="docx",
            format="html",
            outputfile=str(out_path),
            extra_args=extra_args,
        )
        return True, ""
    except Exception as e:
        return False, str(e)


def fallback_htmldocx(html_body: str, title: str, out_path: Path):
    from htmldocx import HtmlToDocx
    doc = Document()
    HtmlToDocx().add_html_to_document(html_body or "<p>(empty)</p>", doc)
    doc.core_properties.title = title or "Converted Document"
    doc.core_properties.language = "en-US"
    apply_language_en_us(doc)
    doc.save(out_path)


def build_docx(title: str, body_html: str, start_level: int, strong_emph: bool) -> bytes:
    if not title:
        title = "Converted Document"

    # Use the body HTML as-is (no injected <h1>)
    html_for_pandoc = body_html or "<p>(empty)</p>"

    with tempfile.TemporaryDirectory() as td:
        out_path = Path(td) / "out.docx"

        ok, err = try_pandoc_convert(
            html_for_pandoc,
            title,
            out_path,
            REF_DOCX if REF_DOCX.exists() else None,
        )

        if not ok:
            # Pandoc failed; fall back (using reference.docx as base if present)
            fallback_htmldocx(body_html, title, out_path)
            doc = Document(str(out_path))
        else:
            doc = Document(str(out_path))

        # 1) Normalize heading levels so the smallest becomes Heading 1
        normalize_heading_levels(doc)

        # 2) Then apply the user-chosen starting level (if > 1)
        if start_level > 1:
            remap_headings(doc, start_level)

        # 3) Map bold/italic to Strong/Emphasis if requested
        if strong_emph:
            bold_italic_to_character_styles(doc, True)

        # 3.5) Add page numbers in the footer for each section
        add_page_numbers(doc)

        # 4) Set language and save to bytes
        doc.core_properties.language = "en-US"
        apply_language_en_us(doc)

        bio = BytesIO()
        doc.save(bio)
        return bio.getvalue()



# --- Submission handling ---
if submitted:
    if not html_body or not html_body.strip():
        st.error("HTML source is required. Please paste some HTML to convert.")
    else:
        try:
            data = build_docx(title, html_body, int(start_level), bool(strong_emph))
            filename = (title or "Converted Document").strip().replace("/", "-") + ".docx"
            st.download_button(
                "Download DOCX file",
                data=data,
                file_name=filename,
                mime=(
                    "application/vnd.openxmlformats-officedocument."
                    "wordprocessingml.document"
                ),
            )
            st.success("Conversion complete. Use the button above to download your file.")
        except Exception as e:
            st.error(f"Conversion failed. Details: {e}")
